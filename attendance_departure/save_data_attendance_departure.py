
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from PIL import Image , ImageTk
from datetime import datetime
import os , sys , shutil
import sqlite3
import time
from tkinter import filedialog
import csv
import docx
from threading import Thread
from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import *
from management_destination.management_exe import copy_resources_folder,manager_Sqlite3


class SaveDataAttendanceDeparture:
    def __init__(self,root):
        self.root = root

        
        self.menubar = Menu(self.root)
        self.file = Menu(self.menubar,tearoff=False)
        self.file.add_command(label='حفظ كـ')
        
        self.menubar.add_cascade(label='حفظ البيانات',command=self.save_file)
        self.root.config(menu=self.menubar)



    def save_file(self):
        """ Save All Data in Docx File """
        self.title = "قائمة الحضور والإنصراف بتأريخ"
        self.date_time = str(datetime.datetime.now().date())

        self.document = docx.Document()
        self.document.add_heading(f"{self.title} - {self.date_time}")

        self.table = self.document.add_table(rows=1, cols=9)
        self.row = self.table.rows[0].cells
        self.table.style = "Colorful List Accent 1"

        self.table_header = [
                      "البريد الإلكتلاوني" , "ملاحضات" , "تاريخ الإنصراف",
                          "تاريخ الحضور" , "الإسم" , "وقت الإنصراف",
                          "وقت الحضور" , "الرقم الوظيفي"

                          ]

        for i in range(8):
            self.table.rows[0].cells[i].text = self.table_header[i]


        self.con = sqlite3.connect(manager_Sqlite3("database_attendance_departure","DataBaseAttendanceDeparture.db"))
        self.curs = self.con.cursor()
        self.res = self.curs.execute("SELECT Email , Notes , date_ansraf , date_7dor ,Name , time_ansraf,time_7dor,ID FROM AttendanceDeparture")
        self.data_sql = self.res.fetchall()

        for self.email , self.notes , self.date_insraf  , self.date_7dor , self.name ,self.time_ansraf , self.time_7dor , self.id_ in self.data_sql:
            self.row = self.table.add_row().cells
            self.row[0].text = self.email
            self.row[1].text = self.notes
            self.row[2].text = str(self.date_insraf)
            self.row[3].text = str(self.date_7dor)
            self.row[4].text = str(self.name)
            self.row[5].text = str(self.time_ansraf)
            self.row[6].text = str(self.time_7dor)
            self.row[7].text = str(self.id_)

        try:
            self.path = filedialog.asksaveasfilename(
                    parent=self.root,
                    initialdir=".", 
                    title="حفظ كـ",
                    defaultextension=".docx",
                    filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*")))

            if self.path:
                if not self.path.lower().endswith(".docx"):
                    self.path +=".docx"

                self.document.save(self.path)
                self.con.commit()

        except Exception:
            pass


