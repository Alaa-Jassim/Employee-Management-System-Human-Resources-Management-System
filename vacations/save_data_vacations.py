from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from PIL import Image , ImageTk
from datetime import datetime
import os , sys , shutil
import sqlite3
import time
from tkinter import filedialog
import csv
import docx
from threading import Thread
from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import * 
from management_destination.management_exe import copy_resources_folder,manager_Sqlite3



class SaveDataVacations:
	def __init__(self,master):
		self.master = master

		self.menubar = Menu(self.master)
		self.file = Menu(self.menubar,tearoff=False)
		self.file.add_command(label='Save As')
		self.menubar.add_cascade(label='حفظ البيانات',command=self.save_vacations_employees)
		self.master.config(menu=self.menubar)


	def save_vacations_employees(self):
		self.title = "قائمة الإجازات بتأريخ"
		self.date_time = str(datetime.datetime.now().date())

		self.doc = docx.Document()
		self.doc.add_heading(f"{self.title} - {self.date_time}")

		self.table = self.doc.add_table(rows=1, cols=9)
		self.row = self.table.rows[0].cells

		self.table.style = 'Colorful List'
		self.table_header = [
							"المصادف يوم"  , "عدد أيام الإجازة" , "القسم الوظيفي" 
							, "إسم الموظف" , "تأريخ الإجازة" , "نوع الإجازة" , "الرقم الوظيفي" , "البريد الإلكتروني"

							]

		for i in range(8):
			self.table.rows[0].cells[i].text = self.table_header[i]

		self.con = sqlite3.connect(manager_Sqlite3("database_vacations","database_vacations.db"))
		self.curs = self.con.cursor()
		
		self.res = self.curs.execute("SELECT coincidence_day ,number_days_vacation , Section, Name,date_coincidence,type_coincidence , ID ,Email FROM VacationsEmployee ")
		self.data_sql = self.res.fetchall()



		for (
			self.coincidence_day , self.number_days_vacation , self.section ,
			self.name , self.date_coincidence , self.type_coincidence ,
			self.id , self.eamil) in self.data_sql:

			self.row = self.table.add_row().cells
			self.row[0].text = self.coincidence_day


			self.row[1].text = self.number_days_vacation
			self.row[2].text = self.section

			self.row[3].text = self.name
			self.row[4].text = str(self.date_coincidence)

			self.row[5].text = self.type_coincidence
			self.row[6].text = str(self.id)
			self.row[7].text = self.eamil


		try:
			self.path = filedialog.asksaveasfilename(
                                            initialdir=".", 
                                            title="Save As", 
                                            defaultextension=".docx",
                                            filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*"))
                                        )

			if self.path:
				if not self.path.lower().endswith(".docx"):
					self.path += ".docx"
				self.doc.save(self.path)
				self.con.commit()

		except Exception as e :
			pass 


