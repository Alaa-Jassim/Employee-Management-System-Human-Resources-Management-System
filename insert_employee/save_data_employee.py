

from tkinter import *
from tkinter import filedialog
import csv
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from PIL import Image , ImageTk
from datetime import datetime
from threading import Thread
import os, sys, shutil
import docx
from datetime import datetime
import sqlite3
from management_destination.management_exe import copy_resources_folder,manager_Sqlite3


class SaveDataInsertEmployee():
	def __init__(self,root):
		self.root = root
		
	
	def save_data(self):
		self.title = "قائمة اسماء الموظفين بتأريخ"
		self.date_time = str(datetime.now().date())

		self.doc = docx.Document()
		self.doc.add_heading(f"{self.title} - {self.date_time}")
		self.table_header = [
			"العمر" , 
			"تأريخ التسجيل",
			"المديمة",
			"القسم",
			"الإجرة الشهرية",
			"الإسم",
			"الرقم الوظيفي",
			"البريد الإلكتروني"
		]
		
		self.table = self.doc.add_table(rows=1, cols=8)
		self.row = self.table.rows[0].cells
		self.table.style = "Colorful List Accent 4"
		
		for i in range(8):
			self.table.rows[0].cells[i].text = self.table_header[i]

		
		self.SQL_QUERY = """ SELECT Age, DateRegistration,
							 City, Section ,
							 Salary, Name ,
							 ID, Email FROM Employees

		 					 """
		
		self.con = sqlite3.connect(manager_Sqlite3("database_employees","DataBaseEmployees.db"))
		self.curs = self.con.cursor()

		self.curs.execute(self.SQL_QUERY)
		self.data_sql = self.curs.fetchall()
		
		for (
			self.age , self.date_registration,
			self.city , self.section ,
			self.salary , self.name ,
			self.id , self.email ) in self.data_sql:
			
			self.row = self.table.add_row().cells
			self.row[0].text = self.age
			self.row[1].text = self.date_registration
			self.row[2].text = self.city 
			self.row[3].text = self.section
			self.row[4].text = self.salary
			self.row[5].text = self.name
			self.row[6].text = str(self.id)
			self.row[7].text = self.email
		
		try:
			self.path = filedialog.asksaveasfilename(
                parent=self.root,
                initialdir=".", 
                title="حفظ كـ", 
                defaultextension=".docx",
                filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*")))
			
			if self.path:
				if not self.path.lower().endswith(".docx"):
					self.path +=".docx"
				self.doc.save(self.path)
				self.con.commit()
		
		except Exception:
			pass 
        