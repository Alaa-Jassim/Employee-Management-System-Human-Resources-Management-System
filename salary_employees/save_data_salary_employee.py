
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from PIL import Image , ImageTk
from datetime import datetime
import os , sys , shutil
import sqlite3
import time
from tkinter import filedialog
import csv
import docx
from threading import Thread
from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import * 
from management_destination.management_exe import copy_resources_folder,manager_Sqlite3

class SaveDataSalaryEmployee:
    def __init__(self,root):
        self.root = root
        
        self.menubar = Menu(self.root)
        self.file = Menu(self.menubar,tearoff=False)
        self.file.add_command(label='حفظ كـ')
        self.menubar.add_cascade(label='حفظ البيانات',command=self.open_file)
        self.root.config(menu=self.menubar)
    

    def open_file(self):
        self.title = "قائمة رواتب الموظفين بتأريخ  "
        self.date_time = str(datetime.datetime.now().date())
        
     
        self.doc = docx.Document()
        self.doc.add_heading(f"{self.title} - {self.date_time}")
        self.table_header = ["إسم الموظف" , "البريد الإلكتروني" , "القسم" , "الإجرة الشهرية" , "الرقم الوظيفي"]
            
            
        self.table = self.doc.add_table(rows=1, cols=5)
        self.row = self.table.rows[0].cells
        self.table.style = "Colorful List Accent 4"

        for i in range(5):
            self.table.rows[0].cells[i].text = self.table_header[i]
                
        self.con = sqlite3.connect(manager_Sqlite3("database_employees","DataBaseEmployees.db"))
        self.curs = self.con.cursor()
        self.res = self.curs.execute('SELECT Name ,Email , Section ,Salary , ID FROM Employees')
        self.data_sql = self.res.fetchall()
            
        for self.name , self.email , self.section , self.salary , self.id in self.data_sql:
            self.row = self.table.add_row().cells
            self.row[0].text = self.name
            self.row[1].text = self.email
            self.row[2].text = self.section
            self.row[3].text = self.salary
            self.row[4].text = str(self.id)
        
        try:
            self.path = filedialog.asksaveasfilename(
                parent=self.root,
                initialdir=".", 
                title="حفظ كـ", 
                defaultextension=".docx",
                filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*")))
            
            if self.path:
                if not self.path.lower().endswith(".docx"):
                    self.path +=".docx"
                self.doc.save(self.path)
                self.con.commit()
        
        except Exception:
            pass 






    

    